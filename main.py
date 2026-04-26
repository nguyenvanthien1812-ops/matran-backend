"""
main.py — FastAPI backend xuất đề thi ra Word (.docx)
Sử dụng python-docx trực tiếp (không qua Pandoc) để kiểm soát
hoàn toàn: bảng rowspan, font Times New Roman, LaTeX → OMML (Word math).

Cấu trúc output Word bám sát 100% exportWord.js:
  Phần I  : Trắc nghiệm nhiều lựa chọn  (câu → A/B/C/D)
  Phần II : Trắc nghiệm Đúng/Sai         (câu → ý a/b/c/d)
  Phần III: Trả lời ngắn                  (câu → ý a/b/c/d)  [tuỳ chọn]
  Phần IV : Tự luận                       (câu → ý a/b/c)
  ------- TRANG MỚI -------
  HƯỚNG DẪN CHẤM:
  Phần I  : Bảng đáp án trắc nghiệm  (Câu | 1 | 2 | ... / Đáp án | A | B | ...)
  Phần II : Bảng Đúng/Sai            (Câu | Ý a | Ý b | Ý c | Ý d)
  Phần III: Bảng Trả lời ngắn        (Câu | Ý a | Ý b | Ý c | Ý d)
  Phần IV : Bảng Tự luận 3 cột       (Câu [rowspan] | Nội dung đáp án | Điểm)
"""

from fastapi import FastAPI, Request
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware

import os, uuid, re, subprocess, tempfile
from copy import deepcopy
from lxml import etree

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────────────────────────────────────
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ══════════════════════════════════════════════════════════════════════════════
#  CONSTANTS
# ══════════════════════════════════════════════════════════════════════════════
FONT_NAME   = "Times New Roman"
SIZE_TITLE  = Pt(14)   # Tiêu đề chính
SIZE_HEAD   = Pt(13)   # Tiêu đề phần
SIZE_BODY   = Pt(12)   # Nội dung câu hỏi / đáp án
SIZE_TABLE  = Pt(11)   # Chữ trong bảng
SIZE_SMALL  = Pt(10)   # Ghi chú nhỏ

BG_HEADER   = "E2E8F0"   # Xanh xám — header bảng
BG_BLUE     = "DBEAFE"   # Xanh nhạt — hàng con TNKQ
BG_GREEN    = "DCFCE7"   # Xanh lá — hàng Tự luận
BG_ORANGE   = "FFEDD5"   # Cam — cột tổng

# ══════════════════════════════════════════════════════════════════════════════
#  UTILS — DỮ LIỆU
# ══════════════════════════════════════════════════════════════════════════════
def get_value(item, keys, default=""):
    if not isinstance(item, dict): return default
    for k in keys:
        val = item.get(k)
        if val is not None and val != "": return val
    return default

def clean_ans(text):
    """Xoá prefix 'Đáp án:' nếu có."""
    if not text: return ""
    return re.sub(r'^(Đáp án:|Đáp án)\s*', '', str(text).strip(), flags=re.IGNORECASE)

def clean_meta(text):
    """Xoá tag nội bộ (Mức độ: ..., Chủ đề: ...) khỏi nội dung câu hỏi."""
    if not text: return ""
    return re.sub(r'\s*\(Mức độ:[^)]*\)\s*', '', str(text)).strip()

def parse_pipe_line(line: str):
    """
    Tách 1 dòng đáp án tự luận thành (noi_dung, diem).
    Hỗ trợ hai format:
      "Nội dung bước || 0.25"          → ưu tiên dấu ||
      "Nội dung bước (0.25 điểm)"      → fallback
    """
    noi_dung = line.strip()
    diem = ""
    pipe_idx = noi_dung.find("||")
    if pipe_idx != -1:
        diem     = noi_dung[pipe_idx + 2:].strip()
        noi_dung = noi_dung[:pipe_idx].strip()
    else:
        m = re.search(r'^(.*?)\s*\(([0-9.,]+)\s*(?:điểm|đ)?\)\s*$', noi_dung)
        if m:
            noi_dung = m.group(1).strip()
            diem     = m.group(2).strip()
    # Xoá dấu * đầu dòng (nếu AI dùng bullet)
    if noi_dung.startswith("*"):
        noi_dung = noi_dung[1:].strip()
    return noi_dung, diem

def get_y_parts(q: dict) -> list:
    """
    Trả về list các ý { label, y, da, di } của câu hỏi.
    Bám sát logic get_y_parts trong main.py gốc & exportWord.js.
    """
    parts = []
    for lbl, key_y, key_da, key_di in [
        ("a", "yA", "dapAnA", "diemA"),
        ("b", "yB", "dapAnB", "diemB"),
        ("c", "yC", "dapAnC", "diemC"),
        ("d", "yD", "dapAnD", "diemD"),
    ]:
        y_val  = str(q.get(key_y,  "") or "").strip()
        da_val = str(q.get(key_da, "") or "").strip()
        di_val = str(q.get(key_di, "") or "").strip()
        if y_val or da_val:
            parts.append({"label": lbl, "y": y_val, "da": da_val, "di": di_val})
    if not parts:
        parts.append({
            "label": "",
            "y": "",
            "da": str(get_value(q, ["answer", "dapAn", "dapAnDung"], "")).strip(),
            "di": str(get_value(q, ["points", "diem"], "")).strip(),
        })
    return parts

# ══════════════════════════════════════════════════════════════════════════════
#  LATEX → OMML  (Word Math Objects)
#  Dùng Pandoc làm converter: Markdown($...$) → .docx → trích OMML
# ══════════════════════════════════════════════════════════════════════════════
_latex_cache: dict = {}

def latex_to_omml(latex_str: str) -> "etree._Element | None":
    """
    Chuyển chuỗi LaTeX → phần tử OMML <m:oMath> để nhúng vào Word.
    Dùng Pandoc: viết file .md tạm → convert → docx → đọc OMML từ document.xml.
    Trả về None nếu thất bại.
    """
    global _latex_cache
    key = latex_str.strip()
    if key in _latex_cache:
        return deepcopy(_latex_cache[key]) if _latex_cache[key] is not None else None

    try:
        # Xác định inline hay block
        if latex_str.strip().startswith("$$"):
            md_text = latex_str.strip()
        else:
            # Đảm bảo bọc $...$
            inner = latex_str.strip().lstrip("$").rstrip("$")
            md_text = f"${inner}$"

        with tempfile.TemporaryDirectory() as tmpdir:
            md_path   = os.path.join(tmpdir, "math.md")
            docx_path = os.path.join(tmpdir, "math.docx")

            with open(md_path, "w", encoding="utf-8") as f:
                f.write(md_text + "\n")

            result = subprocess.run(
                ["pandoc", md_path, "-o", docx_path, "--wrap=none"],
                capture_output=True, text=True, timeout=10
            )
            if result.returncode != 0:
                _latex_cache[key] = None
                return None

            # Đọc document.xml, tìm phần tử m:oMath
            from zipfile import ZipFile
            with ZipFile(docx_path, "r") as z:
                xml_bytes = z.read("word/document.xml")

            tree = etree.fromstring(xml_bytes)
            ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
            omath_list = tree.findall(".//m:oMath", ns)

            if omath_list:
                elem = deepcopy(omath_list[0])
                _latex_cache[key] = elem
                return deepcopy(elem)
            else:
                _latex_cache[key] = None
                return None

    except Exception as e:
        print(f"[LaTeX→OMML] Lỗi: {e} | Input: {latex_str[:80]}")
        _latex_cache[key] = None
        return None

# ══════════════════════════════════════════════════════════════════════════════
#  DOCX HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color: str):
    """Đặt màu nền cho ô bảng một cách an toàn so với giản đồ XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        shd.set(qn("w:fill"), hex_color.upper())
    else:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color.upper())
        tcPr.append(shd)

def add_run_with_font(para, text: str, bold=False, italic=False,
                      font_size=None, font_name=FONT_NAME):
    """Thêm run text vào paragraph với định dạng chuẩn."""
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font_name
    if font_size:
        run.font.size = font_size
    # Đặt font cho East Asian (Hán, Việt)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    return run

def add_cell_rich_text(cell, text: str, bold=False, italic=False,
                       font_size=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    """
    Ghi nội dung rich-text (có thể có LaTeX) vào ô bảng.
    Giữ lại đoạn văn rỗng ban đầu để không làm hỏng cấu trúc tc.
    """
    if font_size is None:
        font_size = SIZE_TABLE

    paragraphs = cell.paragraphs
    if paragraphs:
        current_p = paragraphs[0]
        current_p.text = ""
        current_p.alignment = align
    else:
        current_p = cell.add_paragraph()
        current_p.alignment = align

    parts = re.split(r'(\$\$[\s\S]+?\$\$|\$[^\$]+?\$)', text or "")
    has_latex = any(p.startswith("$") for p in parts if p)

    if not has_latex or not text:
        if text:
            add_run_with_font(current_p, str(text), bold=bold, italic=italic, font_size=font_size)
        return

    first_part = True
    for part in parts:
        if not part: continue
        if part.startswith("$$"):
            if not first_part:
                current_p = cell.add_paragraph()
                current_p.alignment = align
            current_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            omml = latex_to_omml(part)
            if omml is not None: current_p._p.append(omml)
            else: add_run_with_font(current_p, part, italic=True, font_size=font_size)
            current_p = cell.add_paragraph()
            current_p.alignment = align
        elif part.startswith("$"):
            omml = latex_to_omml(part)
            if omml is not None: current_p._p.append(omml)
            else: add_run_with_font(current_p, part, italic=True, font_size=font_size)
        else:
            add_run_with_font(current_p, part, bold=bold, italic=italic, font_size=font_size)
        first_part = False

# ══════════════════════════════════════════════════════════════════════════════
#  XÂY DỰNG HEADER BÌA ĐỀ (bảng 2 cột không viền)
# ══════════════════════════════════════════════════════════════════════════════
def build_header_table(doc, exam_header: dict):
    """Tạo bảng header bìa đề thi (2 cột, không viền)."""
    table = doc.add_table(rows=1, cols=2)
    # Không gán style Table Grid, bảng sẽ không có viền
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cột trái: Sở GD / Trường
    cell_left = table.cell(0, 0)
    p1 = cell_left.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_font(p1, exam_header.get("soGD", "SỞ GIÁO DỤC VÀ ĐÀO TẠO"), font_size=SIZE_BODY)
    p2 = cell_left.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_font(p2, exam_header.get("truong", "TRƯỜNG THPT ..."), bold=True, font_size=SIZE_BODY)

    # Cột phải: Kỳ thi / Môn học / Thời gian / Năm học
    cell_right = table.cell(0, 1)
    for key, is_bold, is_italic in [
        ("kyThi",   True,  False),
        ("monHoc",  True,  False),
        ("thoiGian",False, True ),
        ("namHoc",  False, False),
    ]:
        val = exam_header.get(key, "")
        if not val: continue
        idx = list(cell_right.paragraphs).__len__() - 1
        if idx == 0 and not cell_right.paragraphs[0].text:
            p = cell_right.paragraphs[0]
        else:
            p = cell_right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_font(p, val, bold=is_bold, italic=is_italic, font_size=SIZE_BODY)

    return table


# ══════════════════════════════════════════════════════════════════════════════
#  PHẦN I — TRẮC NGHIỆM NHIỀU LỰA CHỌN
# ══════════════════════════════════════════════════════════════════════════════
def build_section_mcq(doc, questions: list, section_num: int = 1):
    if not questions: return
    title = f"PHẦN {_roman(section_num)}. Câu trắc nghiệm nhiều phương án lựa chọn"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_run_with_font(p, title, bold=True, font_size=SIZE_HEAD)

    for idx, q in enumerate(questions):
        noi_dung = clean_meta(get_value(q, ["noiDung", "noidung", "content"], ""))
        # Số câu in đậm + nội dung
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_run_with_font(p, f"Câu {idx + 1}. ", bold=True, font_size=SIZE_BODY)

        # Xử lý LaTeX trong nội dung
        parts = re.split(r'(\$\$[\s\S]+?\$\$|\$[^\$]+?\$)', noi_dung)
        for part in parts:
            if not part: continue
            if part.startswith("$"):
                omml = latex_to_omml(part)
                if omml is not None: p._p.append(omml)
                else: add_run_with_font(p, part, italic=True, font_size=SIZE_BODY)
            else:
                add_run_with_font(p, part, font_size=SIZE_BODY)

        # Các lựa chọn A/B/C/D — 2 lựa chọn mỗi dòng (giả lập layout ngang)
        labels = ["A", "B", "C", "D"]
        keys   = ["dapAnA", "dapAnB", "dapAnC", "dapAnD"]
        options = [(lbl, str(q.get(k, "") or "")) for lbl, k in zip(labels, keys) if q.get(k, "")]

        # Nếu tất cả ngắn thì xếp 2 cột, còn lại mỗi lựa chọn 1 dòng
        all_short = all(len(v) < 40 for _, v in options)

        if all_short and len(options) == 4:
            # Bảng 4 cột không viền để mô phỏng layout ngang
            tbl = doc.add_table(rows=1, cols=4)
            # Không gán style Table Grid thì bảng sẽ không có viền
            for i, (lbl, val) in enumerate(options):
                cell = tbl.cell(0, i)
                add_cell_rich_text(cell, f"{lbl}. {val}", font_size=SIZE_BODY)
        else:
            for lbl, val in options:
                po = doc.add_paragraph()
                po.paragraph_format.left_indent = Cm(1)
                po.paragraph_format.space_after = Pt(2)
                add_run_with_font(po, f"{lbl}. ", bold=True, font_size=SIZE_BODY)
                parts2 = re.split(r'(\$\$[\s\S]+?\$\$|\$[^\$]+?\$)', val)
                for part in parts2:
                    if not part: continue
                    if part.startswith("$"):
                        omml = latex_to_omml(part)
                        if omml is not None: po._p.append(omml)
                        else: add_run_with_font(po, part, italic=True, font_size=SIZE_BODY)
                    else:
                        add_run_with_font(po, part, font_size=SIZE_BODY)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════════════════════════════════════
#  PHẦN II / III — ĐÚNG/SAI hoặc TRẢ LỜI NGẮN (cấu trúc giống nhau)
# ══════════════════════════════════════════════════════════════════════════════
def build_section_multi_y(doc, questions: list, section_num: int, section_label: str):
    if not questions: return
    title = f"PHẦN {_roman(section_num)}. {section_label}"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_run_with_font(p, title, bold=True, font_size=SIZE_HEAD)

    for idx, q in enumerate(questions):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_run_with_font(p, f"Câu {idx + 1}.", bold=True, font_size=SIZE_BODY)

        y_parts = get_y_parts(q)
        for part in y_parts:
            if not part["y"]: continue
            po = doc.add_paragraph()
            po.paragraph_format.left_indent = Cm(1)
            po.paragraph_format.space_after = Pt(2)
            label_text = f"{part['label']}) " if part["label"] else ""
            y_text = clean_meta(part["y"])
            full = label_text + y_text
            # Xử lý LaTeX
            subs = re.split(r'(\$\$[\s\S]+?\$\$|\$[^\$]+?\$)', full)
            first = True
            for sub in subs:
                if not sub: continue
                if sub.startswith("$"):
                    omml = latex_to_omml(sub)
                    if omml is not None: po._p.append(omml)
                    else: add_run_with_font(po, sub, italic=True, font_size=SIZE_BODY)
                else:
                    add_run_with_font(po, sub, bold=first and bool(label_text), font_size=SIZE_BODY)
                first = False

        doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════════════════════════════════════
#  PHẦN TỰ LUẬN — chỉ hỏi (không có đáp án ở phần đề)
# ══════════════════════════════════════════════════════════════════════════════
def build_section_tu_luan(doc, questions: list, section_num: int):
    if not questions: return
    title = f"PHẦN {_roman(section_num)}. Câu hỏi tự luận"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_run_with_font(p, title, bold=True, font_size=SIZE_HEAD)

    for idx, q in enumerate(questions):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_run_with_font(p, f"Câu {idx + 1}.", bold=True, font_size=SIZE_BODY)

        y_parts = get_y_parts(q)
        for part in y_parts:
            if not part["y"]: continue
            po = doc.add_paragraph()
            po.paragraph_format.left_indent = Cm(1)
            po.paragraph_format.space_after = Pt(4)
            label_text = f"{part['label']}) " if part["label"] else ""
            y_text = clean_meta(part["y"])
            full = label_text + y_text
            subs = re.split(r'(\$\$[\s\S]+?\$\$|\$[^\$]+?\$)', full)
            first = True
            for sub in subs:
                if not sub: continue
                if sub.startswith("$"):
                    omml = latex_to_omml(sub)
                    if omml is not None: po._p.append(omml)
                    else: add_run_with_font(po, sub, italic=True, font_size=SIZE_BODY)
                else:
                    add_run_with_font(po, sub, bold=first and bool(label_text), font_size=SIZE_BODY)
                first = False

        doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
#  HƯỚNG DẪN CHẤM — PHẦN I: BẢNG ĐÁP ÁN TRẮC NGHIỆM
#  2 hàng: [Câu | 1 | 2 | ...] / [Đáp án | A | B | ...]
# ══════════════════════════════════════════════════════════════════════════════
def build_answer_key_mcq(doc, questions: list, diem_moi_cau, section_num: int):
    if not questions: return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_run_with_font(p,
        f"PHẦN {_roman(section_num)}. Trắc nghiệm nhiều phương án lựa chọn "
        f"(Mỗi câu đúng được {diem_moi_cau} điểm)",
        bold=True, font_size=SIZE_HEAD
    )

    n = len(questions)
    table = doc.add_table(rows=2, cols=n + 1)
    table.style = "Table Grid"

    # Hàng 1: Câu | 1 | 2 | ...
    header_row = table.rows[0]
    set_cell_bg(header_row.cells[0], BG_HEADER)
    add_cell_rich_text(header_row.cells[0], "Câu", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=SIZE_TABLE)
    for i in range(n):
        set_cell_bg(header_row.cells[i + 1], BG_HEADER)
        add_cell_rich_text(header_row.cells[i + 1], str(i + 1), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=SIZE_TABLE)

    # Hàng 2: Đáp án | A | B | ...
    ans_row = table.rows[1]
    set_cell_bg(ans_row.cells[0], BG_HEADER)
    add_cell_rich_text(ans_row.cells[0], "Đáp án", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=SIZE_TABLE)
    for i, q in enumerate(questions):
        ans = clean_ans(get_value(q, ["dapAnDung", "answer", "dapAn"], ""))
        add_cell_rich_text(ans_row.cells[i + 1], ans, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=SIZE_TABLE)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
#  HƯỚNG DẪN CHẤM — PHẦN II/III: BẢNG ĐÚNG/SAI hoặc TRẢ LỜI NGẮN
#  Columns: [Câu | Ý a | Ý b | Ý c | Ý d]
# ══════════════════════════════════════════════════════════════════════════════
def build_answer_key_ds_or_tln(doc, questions: list, diem_moi_y, section_num: int, label: str):
    if not questions: return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_run_with_font(p,
        f"PHẦN {_roman(section_num)}. {label} (Mỗi ý đúng được {diem_moi_y} điểm)",
        bold=True, font_size=SIZE_HEAD
    )

    table = doc.add_table(rows=1 + len(questions), cols=5)
    table.style = "Table Grid"

    # Header
    headers = ["Câu", "Ý a", "Ý b", "Ý c", "Ý d"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, BG_HEADER)
        add_cell_rich_text(cell, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=SIZE_TABLE)

    for r_idx, q in enumerate(questions):
        row = table.rows[r_idx + 1]
        # Cột 0: số câu
        set_cell_bg(row.cells[0], BG_HEADER)
        add_cell_rich_text(row.cells[0], str(r_idx + 1), bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER, font_size=SIZE_TABLE)

        # Tách đáp án: "a-Đ, b-S, c-Đ, d-S"  hoặc  "a) 5, b) 12, ..."
        raw = clean_ans(get_value(q, ["dapAnDung", "answer", "dapAn"], ""))
        parts = [x.strip() for x in raw.split(",")]
        for c_idx in range(4):
            val = parts[c_idx] if c_idx < len(parts) else "..."
            add_cell_rich_text(row.cells[c_idx + 1], val,
                               align=WD_ALIGN_PARAGRAPH.CENTER, font_size=SIZE_TABLE)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
#  HƯỚNG DẪN CHẤM — PHẦN TỰ LUẬN
#  Bảng 3 cột: [Câu (rowspan) | Nội dung đáp án | Điểm]
#  Bám sát 100% logic exportWord.js
# ══════════════════════════════════════════════════════════════════════════════
def build_answer_key_tu_luan(doc, questions: list, section_num: int):
    if not questions: return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    add_run_with_font(p,
        f"PHẦN {_roman(section_num)}. Tự luận",
        bold=True, font_size=SIZE_HEAD
    )

    all_rows_data = []  # Danh sách: {"q_idx", "is_title", "text", "diem"}
    q_spans = []        # Lưu khoảng ghép cột Câu cho mỗi câu: (start_row, end_row)

    current_data_row = 1  # Bắt đầu sau header (row 0)

    for q_idx, q in enumerate(questions):
        y_parts_raw = get_y_parts(q)
        # Chỉ lấy ý có nội dung câu hỏi (y)
        y_parts = [p for p in y_parts_raw if p["y"]]
        if not y_parts:
            continue

        start_row_for_q = current_data_row
        rows_for_q = 0

        for part in y_parts:
            # Dòng tiêu đề ý: bold, không có điểm
            lbl_text = f"{part['label']}) " if part["label"] else ""
            all_rows_data.append({
                "q_idx": q_idx,
                "is_title": True,
                "text": lbl_text + clean_meta(part["y"]),
                "diem": ""
            })
            rows_for_q += 1

            # Các dòng bước giải trong đáp án
            if part["da"]:
                lines = [l for l in part["da"].split("\n") if l.strip()]
                for line in lines:
                    nd, diem = parse_pipe_line(line)
                    all_rows_data.append({
                        "q_idx": q_idx,
                        "is_title": False,
                        "text": nd,
                        "diem": diem
                    })
                    rows_for_q += 1

        if rows_for_q > 0:
            q_spans.append((start_row_for_q, start_row_for_q + rows_for_q - 1))
            current_data_row += rows_for_q

    if not all_rows_data:
        return

    # ── XÂY DỰNG BẢNG CHUNG ────────────────────────────────────────────────
    total_data_rows = len(all_rows_data)
    table = doc.add_table(rows=total_data_rows + 1, cols=3)
    table.style = "Table Grid"

    # Header
    hdr_texts = ["Câu", "Nội dung đáp án", "Điểm"]
    hdr_widths = [15, 70, 15]
    for c, (txt, w) in enumerate(zip(hdr_texts, hdr_widths)):
        cell = table.cell(0, c)
        set_cell_bg(cell, BG_HEADER)
        add_cell_rich_text(cell, txt, bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER, font_size=SIZE_TABLE)

    # Dòng dữ liệu
    for r_idx, rd in enumerate(all_rows_data):
        actual_row = r_idx + 1

        # Cột 0: "Câu X", gán tạm vào tất cả các ô, rồi sau đó merge
        cell0 = table.cell(actual_row, 0)
        set_cell_bg(cell0, "F8FAFC")
        add_cell_rich_text(cell0, f"Câu {rd['q_idx'] + 1}", bold=True,
                           align=WD_ALIGN_PARAGRAPH.CENTER, font_size=SIZE_TABLE)

        # Cột 1: nội dung
        cell1 = table.cell(actual_row, 1)
        if rd["is_title"]:
            add_cell_rich_text(cell1, rd["text"], bold=True, font_size=SIZE_TABLE)
            set_cell_bg(cell1, "F1F5F9")
        else:
            add_cell_rich_text(cell1, rd["text"], font_size=SIZE_TABLE)

        # Cột 2: điểm
        cell2 = table.cell(actual_row, 2)
        add_cell_rich_text(cell2, rd["diem"], bold=bool(rd["diem"]),
                           align=WD_ALIGN_PARAGRAPH.CENTER, font_size=SIZE_TABLE)

    # Thực hiện rowspan cột 0 cho từng câu
    for start_r, end_r in q_spans:
        if end_r > start_r:
            cell_start = table.cell(start_r, 0)
            cell_end = table.cell(end_r, 0)
            cell_start.merge(cell_end)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
#  HELPER TỔNG QUÁT
# ══════════════════════════════════════════════════════════════════════════════
def _roman(n: int) -> str:
    mapping = {1:"I",2:"II",3:"III",4:"IV",5:"V",6:"VI"}
    return mapping.get(n, str(n))

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_module.enum.text.WD_BREAK.PAGE)

# ══════════════════════════════════════════════════════════════════════════════
#  ENDPOINT CHÍNH: POST /api/export-docx
# ══════════════════════════════════════════════════════════════════════════════
@app.post("/api/export-docx")
async def export_docx(request: Request):
    import docx as docx_module

    data = await request.json()

    # ── Đọc dữ liệu từ request ─────────────────────────────────────────────
    exam_title  = data.get("examTitle", "ĐỀ KIỂM TRA")
    exam_header = data.get("examHeader", {})
    exam_config = data.get("examConfig", {})

    questions   = data.get("questions")   or data.get("nhieuLuaChon") or []
    dung_sai    = data.get("dung_sai")    or data.get("dungSai")      or []
    tra_loi_ngan= data.get("tra_loi_ngan")or data.get("traLoiNgan")   or []
    tu_luan     = data.get("tu_luan")     or data.get("tuLuan")       or []

    has_tra_loi_ngan = bool(tra_loi_ngan)

    diem_p1 = exam_config.get("diemMoiCauP1", 0.25)
    diem_p2 = exam_config.get("diemMoiYP2",   0.25)
    diem_p3 = exam_config.get("diemMoiYP3",   0.25)

    # ── Tạo Document ───────────────────────────────────────────────────────
    doc = docx_module.Document()

    # Thiết lập trang A4 dọc
    from docx.shared import Cm
    section = doc.sections[0]
    section.page_width   = Cm(21)
    section.page_height  = Cm(29.7)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin   = Cm(2)
    section.bottom_margin= Cm(2)

    # ── HEADER BÌA ─────────────────────────────────────────────────────────
    if exam_header:
        build_header_table(doc, exam_header)
    else:
        p = doc.add_paragraph()
        add_run_with_font(p, "SỞ GIÁO DỤC VÀ ĐÀO TẠO ...", font_size=SIZE_BODY)

    # Tiêu đề đề thi
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    add_run_with_font(p, exam_title, bold=True, font_size=SIZE_TITLE)

    # Thời gian (nếu chưa có trong header)
    if not exam_header:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_font(p2, "Thời gian làm bài: 45 phút", italic=True, font_size=SIZE_BODY)

    doc.add_paragraph()  # Khoảng cách

    # ── XÁC ĐỊNH SỐ THỨ TỰ PHẦN ───────────────────────────────────────────
    sec = 1
    sec_mcq = sec if questions else None;      sec += (1 if questions else 0)
    sec_ds  = sec if dung_sai else None;       sec += (1 if dung_sai else 0)
    sec_tln = sec if tra_loi_ngan else None;   sec += (1 if tra_loi_ngan else 0)
    sec_tl  = sec if tu_luan else None

    # ── PHẦN I — TRẮC NGHIỆM NHIỀU LỰA CHỌN ──────────────────────────────
    if questions and sec_mcq:
        build_section_mcq(doc, questions, sec_mcq)

    # ── PHẦN II — ĐÚNG/SAI ─────────────────────────────────────────────────
    if dung_sai and sec_ds:
        build_section_multi_y(doc, dung_sai, sec_ds, "Câu trắc nghiệm đúng sai")

    # ── PHẦN III — TRẢ LỜI NGẮN ───────────────────────────────────────────
    if tra_loi_ngan and sec_tln:
        build_section_multi_y(doc, tra_loi_ngan, sec_tln, "Câu trắc nghiệm trả lời ngắn")

    # ── PHẦN TỰ LUẬN ──────────────────────────────────────────────────────
    if tu_luan and sec_tl:
        build_section_tu_luan(doc, tu_luan, sec_tl)

    # ══════════════════════════════════════════════════════════════════════
    #  NGẮT TRANG → HƯỚNG DẪN CHẤM
    # ══════════════════════════════════════════════════════════════════════
    # Ngắt trang bằng cách thêm paragraph với run có page break
    pb = doc.add_paragraph()
    pb_run = pb.add_run()
    from docx.oxml import OxmlElement as _OE
    br = _OE("w:br")
    br.set(qn("w:type"), "page")
    pb_run._r.append(br)

    # Tiêu đề HƯỚNG DẪN CHẤM
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(12)
    add_run_with_font(p, "HƯỚNG DẪN CHẤM VÀ BIỂU ĐIỂM", bold=True, font_size=SIZE_TITLE)

    # ── ĐÁP ÁN PHẦN I ─────────────────────────────────────────────────────
    if questions and sec_mcq:
        build_answer_key_mcq(doc, questions, diem_p1, sec_mcq)

    # ── ĐÁP ÁN PHẦN II — ĐÚNG/SAI ─────────────────────────────────────────
    if dung_sai and sec_ds:
        build_answer_key_ds_or_tln(
            doc, dung_sai, diem_p2, sec_ds, "Trắc nghiệm Đúng/Sai"
        )

    # ── ĐÁP ÁN PHẦN III — TRẢ LỜI NGẮN ────────────────────────────────────
    if tra_loi_ngan and sec_tln:
        build_answer_key_ds_or_tln(
            doc, tra_loi_ngan, diem_p3, sec_tln, "Trắc nghiệm Trả lời ngắn"
        )

    # ── ĐÁP ÁN PHẦN TỰ LUẬN ───────────────────────────────────────────────
    if tu_luan and sec_tl:
        build_answer_key_tu_luan(doc, tu_luan, sec_tl)

    # ── LƯU FILE & TRẢ VỀ ─────────────────────────────────────────────────
    filename = f"/tmp/export_exam_{uuid.uuid4().hex}.docx"
    try:
        doc.save(filename)
    except Exception as e:
        return {"error": f"Lỗi khi lưu file: {str(e)}"}

    return FileResponse(
        filename,
        filename="De_Thi_Chuẩn.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=None,
    )


# ══════════════════════════════════════════════════════════════════════════════
#  ENDPOINT: POST /api/generate-graph  (Matplotlib đồ thị)
# ══════════════════════════════════════════════════════════════════════════════
@app.post("/api/generate-graph")
async def generate_graph_endpoint(request: Request):
    """
    Nhận JSON metadata đồ thị, trả về PNG base64 + file path.
    Body: { "loai": "do_thi_ham_so", "hamSo": "x**2", "xRange": [-5,5], ... }
    """
    from backend.graph_service import generate_graph
    data = await request.json()
    result = generate_graph(data)
    return result


# ══════════════════════════════════════════════════════════════════════════════
#  ENDPOINT: POST /api/export-docx-with-images  (Word + hình ảnh Matplotlib)
#  Mở rộng từ /api/export-docx, thêm khả năng chèn PNG từ Matplotlib
# ══════════════════════════════════════════════════════════════════════════════
@app.post("/api/export-docx-with-images")
async def export_docx_with_images(request: Request):
    """
    Giống /api/export-docx nhưng hỗ trợ thêm field 'hinhAnh' trong mỗi câu hỏi.
    Nếu câu hỏi có hinhAnh metadata → gọi Matplotlib vẽ → chèn PNG vào Word.
    """
    import docx as docx_module
    from backend.graph_service import generate_graph

    data = await request.json()

    # ── Đọc dữ liệu từ request (giống hệt /api/export-docx) ────────────
    exam_title  = data.get("examTitle", "ĐỀ KIỂM TRA")
    exam_header = data.get("examHeader", {})
    exam_config = data.get("examConfig", {})

    questions   = data.get("questions")   or data.get("nhieuLuaChon") or []
    dung_sai    = data.get("dung_sai")    or data.get("dungSai")      or []
    tra_loi_ngan= data.get("tra_loi_ngan")or data.get("traLoiNgan")   or []
    tu_luan     = data.get("tu_luan")     or data.get("tuLuan")       or []

    diem_p1 = exam_config.get("diemMoiCauP1", 0.25)
    diem_p2 = exam_config.get("diemMoiYP2",   0.25)
    diem_p3 = exam_config.get("diemMoiYP3",   0.25)

    # ── Tạo Document ───────────────────────────────────────────────────
    doc = docx_module.Document()

    from docx.shared import Cm
    section = doc.sections[0]
    section.page_width   = Cm(21)
    section.page_height  = Cm(29.7)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin   = Cm(2)
    section.bottom_margin= Cm(2)

    # ── HEADER BÌA ─────────────────────────────────────────────────────
    if exam_header:
        build_header_table(doc, exam_header)

    # Tiêu đề đề thi
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    add_run_with_font(p, exam_title, bold=True, font_size=SIZE_TITLE)

    doc.add_paragraph()

    # ── Helper: chèn hình Matplotlib vào doc nếu có ────────────────────
    def _insert_graph_if_needed(q: dict):
        """Nếu câu hỏi có field 'hinhAnh', gọi Matplotlib vẽ và chèn PNG."""
        hinh_anh = q.get("hinhAnh")
        if not hinh_anh or not isinstance(hinh_anh, dict):
            return
        result = generate_graph(hinh_anh)
        if result.get("success") and result.get("filePath"):
            try:
                doc.add_picture(result["filePath"], width=Inches(4))
                # Căn giữa hình
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                last_paragraph.paragraph_format.space_after = Pt(8)
            except Exception as e:
                p_err = doc.add_paragraph()
                add_run_with_font(p_err, f"[Lỗi chèn hình: {e}]",
                                  italic=True, font_size=SIZE_SMALL)

    # ── Xác định số thứ tự phần ────────────────────────────────────────
    sec = 1
    sec_mcq = sec if questions else None;      sec += (1 if questions else 0)
    sec_ds  = sec if dung_sai else None;       sec += (1 if dung_sai else 0)
    sec_tln = sec if tra_loi_ngan else None;   sec += (1 if tra_loi_ngan else 0)
    sec_tl  = sec if tu_luan else None

    # ── PHẦN I — TRẮC NGHIỆM ──────────────────────────────────────────
    if questions and sec_mcq:
        build_section_mcq(doc, questions, sec_mcq)
        # Chèn hình cho từng câu MCQ
        for q in questions:
            _insert_graph_if_needed(q)

    # ── PHẦN II — ĐÚNG/SAI ─────────────────────────────────────────────
    if dung_sai and sec_ds:
        build_section_multi_y(doc, dung_sai, sec_ds, "Câu trắc nghiệm đúng sai")
        for q in dung_sai:
            _insert_graph_if_needed(q)

    # ── PHẦN III — TRẢ LỜI NGẮN ───────────────────────────────────────
    if tra_loi_ngan and sec_tln:
        build_section_multi_y(doc, tra_loi_ngan, sec_tln,
                              "Câu trắc nghiệm trả lời ngắn")
        for q in tra_loi_ngan:
            _insert_graph_if_needed(q)

    # ── PHẦN TỰ LUẬN ──────────────────────────────────────────────────
    if tu_luan and sec_tl:
        build_section_tu_luan(doc, tu_luan, sec_tl)
        for q in tu_luan:
            _insert_graph_if_needed(q)

    # ── HƯỚNG DẪN CHẤM ────────────────────────────────────────────────
    pb = doc.add_paragraph()
    pb_run = pb.add_run()
    from docx.oxml import OxmlElement as _OE2
    br2 = _OE2("w:br")
    br2.set(qn("w:type"), "page")
    pb_run._r.append(br2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(12)
    add_run_with_font(p, "HƯỚNG DẪN CHẤM VÀ BIỂU ĐIỂM", bold=True,
                      font_size=SIZE_TITLE)

    if questions and sec_mcq:
        build_answer_key_mcq(doc, questions, diem_p1, sec_mcq)
    if dung_sai and sec_ds:
        build_answer_key_ds_or_tln(doc, dung_sai, diem_p2, sec_ds,
                                    "Trắc nghiệm Đúng/Sai")
    if tra_loi_ngan and sec_tln:
        build_answer_key_ds_or_tln(doc, tra_loi_ngan, diem_p3, sec_tln,
                                    "Trắc nghiệm Trả lời ngắn")
    if tu_luan and sec_tl:
        build_answer_key_tu_luan(doc, tu_luan, sec_tl)

    # ── LƯU FILE & TRẢ VỀ ─────────────────────────────────────────────
    filename = f"/tmp/export_exam_img_{uuid.uuid4().hex}.docx"
    try:
        doc.save(filename)
    except Exception as e:
        return {"error": f"Lỗi khi lưu file: {str(e)}"}

    return FileResponse(
        filename,
        filename="De_Thi_Co_Hinh.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=None,
    )


# ══════════════════════════════════════════════════════════════════════════════
#  ENDPOINT: GET /api/health  (Kiểm tra backend đang chạy)
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/api/health")
async def health_check():
    return {"status": "ok", "matplotlib": True, "pandoc": True}


# ══════════════════════════════════════════════════════════════════════════════
#  CHẠY TRỰC TIẾP (DEVELOPMENT)
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)